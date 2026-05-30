from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "share_widget_figma_prototype_guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E2EC", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(9.5)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths, header_fill="E8EEF5"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        add_cell_text(cell, h, bold=True, color="0B2545")
        set_cell_shading(cell, header_fill)
        set_cell_border(cell)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            add_cell_text(cells[i], value)
            set_cell_border(cells[i])
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, fill="F2FBF8", border="B7E4D8"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=border, size="10")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string("126B5B")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Malgun Gothic"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor.from_string("1F2937")
    doc.add_paragraph()


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10)


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    return p


def add_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        rest = text[len(bold_prefix):]
        r2 = p.add_run(rest)
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for r in runs:
        r.font.name = "Malgun Gothic"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        r.font.size = Pt(10.5)
    return p


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 24, "0B2545", 0, 10),
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("공유 위젯 앱 프로토타입 Figma 구현 가이드")
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("667085")


doc = Document()
style_document(doc)

title = doc.add_paragraph(style="Title")
title.add_run("공유 위젯 기반 정보 검증 앱 프로토타입 Figma 구현 가이드")
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run("프로토타입 1: 앱 진입 없이 공유 화면에서 검증 / 프로토타입 2: 앱 안에서 기록 관리와 정정 재발신")
subtitle_run.font.name = "Malgun Gothic"
subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
subtitle_run.font.size = Pt(11)
subtitle_run.font.color.rgb = RGBColor.from_string("475467")

add_callout(
    doc,
    "프로토타입 공통 예시 콘텐츠",
    "URL: https://youtu.be/pAxPZHBMGF0?si=hw2a6emoLlcGuulQ\n제목: [건강 정보] 막걸리 효능 제대로 보기 위해선 무조건 이렇게 드세요\n출처 표기: 유튜브에서 가져옴 / 채널명 예시: 건강힐링 / 주제: 건강 정보",
    fill="F7FAFF",
    border="B8D4F6",
)

add_h(doc, "1. 문서 목적과 제작 범위", 1)
add_p(doc, "이 문서는 처음 보는 Figma 작업자가 글만 보고 앱 프로토타입을 구현할 수 있도록 만든 화면 설계 가이드입니다. 실제 개발 명세가 아니라 발표와 사용자 테스트를 위한 클릭형 프로토타입 명세입니다.")
add_bullet(doc, "앱 이름은 미정으로 두고, Figma에서는 임시명으로 '바른공유'를 사용합니다. 최종 명칭이 생기면 전체 교체합니다.")
add_bullet(doc, "체크리스트 명칭은 확정값인 '바른정보길잡이'를 사용합니다.")
add_bullet(doc, "지수명은 미정이므로 문서와 화면에서는 '검증지수(가칭)'로 표기합니다.")
add_bullet(doc, "설정, 튜토리얼, 최근 사용 앱, 미디어 리터러시 교육 버튼은 프로토타입 설정상 남기되 일부는 가짜 버튼으로 처리합니다.")

add_h(doc, "2. 핵심 사용자 경험", 1)
add_table(
    doc,
    ["구분", "사용 상황", "핵심 흐름", "Figma에서 보여줄 결과"],
    [
        ["프로토타입 1", "유튜브 등 외부 앱에서 공유 버튼을 누른 직후", "공유 시트 → 정보 검증해서 공유하기 → 체크리스트/참고자료 → 대상 선택 → 메시지 입력 → 공유 완료", "앱에 들어가지 않아도 검증을 거쳐 공유할 수 있음을 보여줌"],
        ["프로토타입 2", "앱을 내려받아 앱 안에서 기록을 관리하는 상황", "홈/내 서랍 → 공유 기록 확인 → 검증 다시 하기 또는 정정 → 이전 수신자에게 정정 메시지 재발신", "과거 공유 정보를 관리하고 정정하는 서비스 가치를 보여줌"],
    ],
    [1.25, 1.8, 2.1, 1.35],
)

add_h(doc, "3. Figma 파일 기본 세팅", 1)
add_number(doc, "페이지를 4개 만듭니다: 00_Cover, 01_Design System, 02_Proto 1_Widget Share, 03_Proto 2_App.")
add_number(doc, "모바일 프레임은 iPhone 14 기준 390 x 844 px를 기본으로 사용합니다. 첨부 이미지처럼 상단 상태바와 하단 홈 인디케이터를 포함합니다.")
add_number(doc, "모든 주요 터치 영역은 최소 48 x 48 px로 만듭니다. 고령 사용자도 누르기 쉽게 버튼 높이는 56 px 이상을 권장합니다.")
add_number(doc, "본문 폰트는 Pretendard 또는 Noto Sans KR을 사용합니다. 제목 22-26 px, 섹션 제목 20-22 px, 본문 16-18 px, 보조문구 14-15 px로 잡습니다.")
add_number(doc, "Auto Layout을 적극 사용합니다. 카드 내부 여백 16 px, 카드 간격 12-16 px, 화면 좌우 여백 20 px를 기본값으로 둡니다.")

add_h(doc, "4. 시각 디자인 원칙", 1)
add_table(
    doc,
    ["항목", "권장값", "이유"],
    [
        ["주요 색", "검증 액션: 짙은 청록 #1E7F7A / 공유 완료: 파랑 #2457A7", "신뢰, 확인, 공공성을 느끼게 함"],
        ["검증한 정보", "초록 계열 배지 #2EAD73", "이미 검증 절차를 거친 항목이라는 즉시 인지"],
        ["미검증 정보", "빨강 계열 배지 #D92D20", "그냥 보낸 정보와 검증한 정보를 강하게 구분"],
        ["검증 결과", "상 #157347, 중 #B7791F, 하 #C0392B", "초록 안에서도 믿을 만한 정도를 구분"],
        ["카드 형태", "반경 16 px, 그림자 약하게, 배경 흰색", "첨부 예시의 카드 기반 모바일 UI를 정돈된 형태로 재해석"],
        ["톤", "큰 글자, 짧은 문장, 한 화면 한 과업", "처음 쓰는 사람과 고령 사용자가 흐름을 놓치지 않도록 함"],
    ],
    [1.2, 2.1, 3.2],
)

add_h(doc, "5. 공통 컴포넌트", 1)
add_table(
    doc,
    ["컴포넌트", "구성", "상태/변형"],
    [
        ["콘텐츠 미리보기 카드", "썸네일 72 x 72, 제목 2줄, 출처, URL 숨김 또는 작은 글씨", "기본/로딩/링크 오류"],
        ["바른정보길잡이 항목", "체크박스, 질문 텍스트, 도움 버튼", "미체크/체크됨/도움 열림"],
        ["도움 버튼", "작은 외곽선 버튼, 라벨 '도움'", "누르면 하단 시트 또는 카드 확장"],
        ["참고자료 입력", "검색 입력창, AI 도움받기 버튼, 빈 상태 카드", "빈 상태/검색 중/자료 추가됨"],
        ["검증지수 표시", "상/중/하 라벨, 0-100 점수, 진행바", "미검증/하/중/상"],
        ["오늘의 검증 도장", "원형 도장 영역, 칭찬 문구", "비어 있음/도장 찍힘"],
        ["공유 대상 칩", "친구, 가족, 단톡방, 새 그룹", "선택 전/선택됨"],
        ["기록 카드", "썸네일, 제목, 메시지, 송수신자, 검증상태, CTA", "검증한 정보/미검증 정보"],
    ],
    [1.55, 2.75, 2.2],
)

add_h(doc, "6. 바른정보길잡이 체크리스트 5개", 1)
add_p(doc, "체크리스트는 '예'라고 답할 수 있으면 체크하는 구조입니다. 문구는 판단이 쉬운 생활 언어로 쓰고, 한 항목은 두 줄을 넘기지 않습니다.")
add_table(
    doc,
    ["번호", "체크리스트 문구", "도움 버튼을 눌렀을 때"],
    [
        ["1", "정보의 출처가 누구인지 확인했나요?", "채널명, 작성자, 기관명, 영상 설명란을 확인하라고 안내합니다."],
        ["2", "같은 내용을 믿을 만한 다른 자료에서도 확인했나요?", "정부/병원/언론/학회 등 참고자료 검색을 제안합니다."],
        ["3", "건강, 돈, 구매를 강하게 권하는 표현이 있는지 살펴봤나요?", "건강 정보와 구매 유도 표현은 특히 주의해서 보라고 안내합니다."],
        ["4", "제목이나 썸네일이 너무 자극적이지 않은지 확인했나요?", "'무조건', '싹 사라집니다' 같은 과장 표현을 예로 보여줍니다."],
        ["5", "이전에 보냈던 정보에 대한 정정인가요?", "체크 후 도움을 누르면 내 서랍 아카이브 팝업을 열어 과거 공유 목록을 보여줍니다."],
    ],
    [0.55, 2.85, 3.1],
)

add_h(doc, "7. 검증지수(가칭) 산출 방식", 1)
add_p(doc, "프로토타입에서는 실제 계산처럼 보이도록 단순한 룰을 화면에 반영합니다. 점수는 체크리스트, 참고자료 개수, 참고자료 출처 신뢰도를 합산해 자동 산출되는 설정입니다.")
add_table(
    doc,
    ["요소", "점수", "화면 반영"],
    [
        ["체크리스트 체크", "항목당 12점, 최대 60점", "체크할 때마다 진행바가 올라감"],
        ["참고자료 개수", "1개 10점, 2개 이상 20점", "참고자료 카드가 추가되면 점수 상승"],
        ["출처 신뢰도", "공식기관/병원/학회/공영자료 20점, 일반 블로그 5점", "자료 옆에 '공식자료' 또는 '일반자료' 배지"],
        ["주의 요소", "과장 제목, 구매 유도, 건강 단정 표현 발견 시 -10점", "도움 메시지에 주의 배지 표시"],
    ],
    [1.4, 1.55, 3.55],
)
add_p(doc, "등급 기준: 상 80-100점, 중 50-79점, 하 0-49점. 검증 절차를 전혀 거치지 않은 공유 기록은 점수가 아니라 빨간색 '미검증 정보'로 먼저 표시합니다.")

add_h(doc, "8. 프로토타입 1: 공유 위젯에서 검증하고 공유하기", 1)
add_p(doc, "목표: 앱을 열지 않고도 외부 공유 화면에서 검증 과정을 시작할 수 있음을 보여줍니다. 첨부된 유튜브 공유 화면 예시를 기준으로, 기본 공유 시트 아래에 '정보 검증해서 공유하기' CTA를 추가합니다.")
add_table(
    doc,
    ["프레임명", "화면 구성", "사용자 액션", "다음 연결"],
    [
        ["P1-01 유튜브 공유 화면", "상단 유튜브 영상 카드, 공유 대상 아이콘, 하단 큰 CTA '정보 검증해서 공유하기'", "CTA 탭", "P1-02 정보 검증 화면"],
        ["P1-02 정보 검증 화면", "상단 앱바, 콘텐츠 미리보기 카드, 바른정보길잡이 5개, 참고자료 영역, 하단 고정 버튼", "체크박스 선택 또는 도움 탭", "도움 시트 또는 점수 상승"],
        ["P1-03 도움 시트", "선택 항목 설명, 주의 문구, 추천 확인 방법", "확인했어요 탭", "P1-02로 돌아가 체크 상태 반영"],
        ["P1-04 참고자료 AI 도움", "검색 중 상태 후 참고자료 2개 카드 표시", "자료 추가 탭", "P1-02 점수 상승"],
        ["P1-05 공유 대상 고르기", "친구 목록, 가족 프리셋, 새 그룹 만들기", "대상 선택", "P1-06 메시지 입력"],
        ["P1-06 메시지 입력", "자동 문구, 사용자 추가 메시지, 검증지수 미리보기", "공유하기 탭", "P1-07 공유 완료"],
        ["P1-07 공유 완료", "완료 체크, 도장 획득, 공유 대상 요약", "내 서랍 보기 또는 닫기", "프로토 종료"],
    ],
    [1.35, 2.5, 1.45, 1.2],
)

add_h(doc, "9. 프로토타입 1 화면별 세부 명세", 2)
add_h(doc, "P1-01 유튜브 공유 화면", 3)
add_bullet(doc, "영상 제목은 '[건강 정보] 막걸리 효능 제대로 보기 위해선 무조건 이렇게 드세요'로 넣습니다.")
add_bullet(doc, "썸네일은 첨부된 막걸리 영상 캡처를 사용합니다. Figma에서는 이미지 위에 17:03 재생시간 배지를 검은 반투명 박스로 배치합니다.")
add_bullet(doc, "기존 공유 아이콘 아래에 앱 위젯 CTA를 추가합니다. 버튼 문구: '정보 검증해서 공유하기'. 왼쪽에는 방패 또는 체크 아이콘을 둡니다.")

add_h(doc, "P1-02 정보 검증 화면", 3)
add_bullet(doc, "앱바: 뒤로가기, 제목 '정보 검증하기', 설정 아이콘. 설정 아이콘은 가짜 버튼으로 둡니다.")
add_bullet(doc, "콘텐츠 미리보기 카드: 썸네일, 제목 2줄, '유튜브에서 가져옴', 링크 복사 완료 상태를 보여줍니다.")
add_bullet(doc, "섹션 제목은 '바른정보길잡이'로 표기합니다. 사용자가 체크할 때마다 하단 검증지수 진행바가 0 → 35 → 60 → 80처럼 상승합니다.")
add_bullet(doc, "참고자료 영역의 빈 상태 문구: '관련 자료를 검색하거나 AI에게 도움을 요청해보세요'.")
add_bullet(doc, "하단 고정 버튼은 기본 '다음: 공유 대상 고르기'입니다. 검증지수가 하인 경우에도 이동은 가능하지만 주의 안내 모달을 한 번 보여줍니다.")

add_h(doc, "P1-03 도움 시트 예시 문구", 3)
add_table(
    doc,
    ["항목", "도움 문구"],
    [
        ["출처 확인", "영상 설명란과 채널 소개를 확인해보세요. 병원, 학회, 정부기관처럼 책임 있는 출처인지 살펴보면 좋아요."],
        ["다른 자료 확인", "같은 내용을 질병관리청, 식약처, 병원 건강정보, 공영 교육 자료에서도 찾을 수 있는지 확인해보세요."],
        ["건강/구매 주의", "이 영상은 건강 정보입니다. 특정 음식이나 제품이 병을 고친다고 단정하면 주의해서 시청해주세요."],
        ["자극적 표현", "'무조건', '싹 사라집니다', '필수' 같은 표현은 과장일 수 있어요. 본문 근거를 함께 확인하세요."],
        ["정정 여부", "이전에 보낸 정보와 반대되는 내용이라면 내 서랍에서 예전 공유 기록을 찾아 정정 메시지로 다시 보낼 수 있어요."],
    ],
    [1.4, 5.1],
)

add_h(doc, "P1-06 메시지 입력 기본 문구", 3)
add_callout(
    doc,
    "공유 메시지 예시",
    "막걸리 건강 정보 영상이라서 바른정보길잡이로 한 번 확인해봤어요.\n검증지수(가칭): 중\n건강 관련 내용은 사람마다 다를 수 있으니 참고용으로만 봐주세요.\n\n링크: https://youtu.be/pAxPZHBMGF0?si=hw2a6emoLlcGuulQ",
)

add_h(doc, "10. 프로토타입 2: 앱 안에서 기록 관리와 정정 재발신", 1)
add_p(doc, "목표: 사용자가 앱에 들어온 뒤 내 서랍에서 이전 공유 기록을 확인하고, 잘못된 정보 또는 반대 정보를 발견했을 때 같은 수신자에게 정정 메시지를 다시 보내는 과정을 보여줍니다.")
add_table(
    doc,
    ["프레임명", "화면 구성", "사용자 액션", "다음 연결"],
    [
        ["P2-01 홈", "프로필, 오늘의 검증 도장, 최근 공유 기록 2개, 미디어 리터러시 영상 바로가기, 하단 탭", "내 서랍 탭", "P2-02 내 서랍"],
        ["P2-02 내 서랍 목록", "주제별/검증도별 필터, 검증한 정보 초록 카드, 미검증 정보 빨간 카드", "막걸리 기록 선택", "P2-03 기록 상세"],
        ["P2-03 기록 상세", "썸네일, 제목, 보낸 메시지, 송수신자, 검증지수, 참고자료", "정정 버튼 탭", "P2-04 정정 정보 입력"],
        ["P2-04 정정 정보 입력", "새 링크 입력, 자동 미리보기, 정정 사유 입력, 기존 수신자 자동 선택", "정정 메시지 만들기", "P2-05 정정 메시지 확인"],
        ["P2-05 정정 메시지 확인", "자동 생성 문구와 공유 대상 목록", "재발신하기 탭", "P2-06 정정 완료"],
        ["P2-06 정정 완료", "완료 안내, 오늘의 검증 도장 획득, 기록 카드 상태 변경", "내 서랍으로", "P2-02로 돌아감"],
    ],
    [1.35, 2.45, 1.45, 1.25],
)

add_h(doc, "11. 프로토타입 2 화면별 세부 명세", 2)
add_h(doc, "P2-01 홈", 3)
add_bullet(doc, "상단 프로필 영역: 프로필 사진 자리, '나의 기본정보', '기본 연락 플랫폼', '나의 검증지수(가칭)'를 보여줍니다.")
add_bullet(doc, "오늘의 검증 카드: 처음에는 빈 도장 원이 있고, 검증 완료 후 '칭찬해요! 오늘의 검증 도장' 상태로 바뀝니다. 프로토타입에서는 도장 1종만 구현하고, 설정상 다양한 도장을 모을 수 있다고 둡니다.")
add_bullet(doc, "체크리스트 사용자 맞춤화: '바른정보길잡이 바꾸기' 버튼을 둡니다. 클릭하면 편집 화면으로 이동하되, 상세 편집은 가짜 버튼 또는 간단 화면으로 처리합니다.")
add_bullet(doc, "미디어 리터러시 영상 바로가기: 공식 교육 웹사이트 또는 미디온 플러스 영상으로 이동하는 설정의 가짜 버튼입니다.")

add_h(doc, "P2-02 내 서랍 목록", 3)
add_bullet(doc, "상단 필터: 주제별 칩(건강, 정치, 금융, 생활), 검증도별 칩(상, 중, 하, 미검증)을 둡니다.")
add_bullet(doc, "검증한 정보 카드는 초록색 띠와 '검증한 정보' 배지를 붙입니다. 버튼은 '검증 다시 하기'로 표시합니다.")
add_bullet(doc, "미검증 정보 카드는 빨간색 띠와 '미검증 정보' 배지를 붙입니다. 버튼은 '검증하기'로 표시합니다.")
add_bullet(doc, "막걸리 영상 기록 카드는 제목, 썸네일, 보낸 메시지 일부, 보낸 사람/받은 사람, 검증지수 '중'을 보여줍니다.")

add_h(doc, "P2-04 정정 정보 입력", 3)
add_bullet(doc, "사용자가 '정정'을 누르면 새로 찾은 반대 정보 또는 정정 링크를 입력하는 화면을 보여줍니다.")
add_bullet(doc, "이전 수신자 목록은 자동으로 선택되어 있습니다. 사용자는 체크를 해제할 수 있습니다.")
add_bullet(doc, "자동 첨부 문구는 다음 화면에서 미리 보여줍니다.")

add_h(doc, "P2-05 자동 정정 메시지", 3)
add_callout(
    doc,
    "정정 메시지 예시",
    "이전에 보낸 '[건강 정보] 막걸리 효능 제대로 보기 위해선 무조건 이렇게 드세요' 정보와 반대되는 내용을 찾았어요.\n건강 정보는 사람마다 다르게 적용될 수 있어서, 아래 정정 내용을 함께 확인해주세요.\n\n정정 정보: 사용자가 새로 입력한 링크 또는 설명\n기존 공유 대상: 가족 단톡방, 조한결, 허윤",
    fill="FFF8E6",
    border="F0C36D",
)

add_h(doc, "12. 아카이브 팝업: 이전에 보낸 정보에 대한 정정", 1)
add_p(doc, "프로토타입 1에서도 체크리스트 5번 '이전에 보냈던 정보에 대한 정정인가요?'를 체크하고 도움 버튼을 누르면 내 서랍의 공유 기록 목록이 팝업으로 뜹니다.")
add_table(
    doc,
    ["영역", "내용"],
    [
        ["팝업 제목", "정정할 이전 정보를 골라주세요"],
        ["목록 카드", "썸네일, 제목, 보낸 날짜, 수신자, 검증상태"],
        ["선택 동작", "카드 선택 시 하단에 '이 정보에 대한 정정으로 보내기' 버튼 활성화"],
        ["자동 문구", "이전에 보낸 A정보에 반대되는 정보를 찾았어요."],
        ["연결", "공유 대상은 기존 수신자로 자동 설정되고, 사용자는 메시지를 추가할 수 있음"],
    ],
    [1.45, 5.05],
)

add_h(doc, "13. Figma 프로토타입 연결 규칙", 1)
add_number(doc, "주요 버튼은 모두 Tap 인터랙션으로 연결합니다. 전환은 Smart Animate 또는 Move In 300ms를 사용합니다.")
add_number(doc, "도움 버튼, 참고자료 검색, 정정 아카이브는 Overlay로 띄웁니다. 배경은 #000000 30% 투명도 dim을 적용합니다.")
add_number(doc, "체크박스는 컴포넌트 variant로 만듭니다. 클릭 시 미체크 → 체크됨 variant로 바꾸고 검증지수 바가 올라간 다음 프레임으로 연결합니다.")
add_number(doc, "검증지수는 0, 35, 60, 80 세 가지 이상 상태 프레임을 만들어 진행감을 보여줍니다.")
add_number(doc, "하단 탭은 홈, 도움말, 내 서랍 3개로 구성합니다. 홈과 내 서랍만 실제 연결하고 도움말은 가짜 버튼으로 둡니다.")

add_h(doc, "14. 화면에 넣을 고정 문구", 1)
add_table(
    doc,
    ["위치", "문구"],
    [
        ["공유 위젯 CTA", "정보 검증해서 공유하기"],
        ["검증 화면 제목", "정보 검증하기"],
        ["체크리스트 제목", "바른정보길잡이"],
        ["참고자료 빈 상태", "관련 자료를 검색하거나 AI에게 도움을 요청해보세요"],
        ["AI 버튼", "AI 도움받기"],
        ["다음 버튼", "다음: 공유 대상 고르기"],
        ["공유 완료", "검증한 정보를 공유했어요"],
        ["도장 획득", "칭찬해요! 오늘의 검증 도장을 받았어요"],
        ["내 서랍", "내가 보낸 기록 모아보기"],
        ["정정 CTA", "정정해서 다시 보내기"],
    ],
    [1.75, 4.75],
)

add_h(doc, "15. 사용성 체크리스트", 1)
add_bullet(doc, "한 화면에서 사용자가 해야 할 일이 하나만 보이는지 확인합니다.")
add_bullet(doc, "체크박스, 도움 버튼, 다음 버튼의 터치 영역이 충분히 큰지 확인합니다.")
add_bullet(doc, "검증한 정보와 미검증 정보가 색과 문구로 모두 구분되는지 확인합니다.")
add_bullet(doc, "건강 정보에 대한 주의 문구가 공유 전 메시지에 자동 포함되는지 확인합니다.")
add_bullet(doc, "정정 흐름에서 기존 수신자가 자동 선택되어 '누구에게 다시 보내야 하는지' 고민하지 않아도 되는지 확인합니다.")
add_bullet(doc, "설정, 튜토리얼, 최근 사용 앱, 미디어 리터러시 교육은 프로토타입에서 눌렀을 때 막히지 않도록 토스트 또는 준비중 오버레이를 연결합니다.")

add_h(doc, "16. 제작자가 마지막에 확인할 것", 1)
add_number(doc, "프로토타입 1은 외부 공유 화면에서 시작해 공유 완료까지 끊기지 않아야 합니다.")
add_number(doc, "프로토타입 2는 홈에서 내 서랍으로 들어가 정정 재발신 완료까지 끊기지 않아야 합니다.")
add_number(doc, "예시 콘텐츠의 URL과 제목이 모든 화면에서 동일해야 합니다.")
add_number(doc, "바른정보길잡이 5개 항목과 검증지수 산출 방식이 화면 설명과 일치해야 합니다.")
add_number(doc, "발표자가 설명하지 않아도 버튼 문구만 보고 다음 행동을 이해할 수 있어야 합니다.")

doc.save(OUT)
print(OUT)
